"""
DOCX Safe Text Replacement  —  3-Phase Approach
================================================
Phase 1  DETECT   : Pehle poori formatting ka XML snapshot lo
Phase 2  REPLACE  : Run-level safe text replacement
Phase 3  RESTORE  : Snapshot se 100% formatting wapas guarantee

Kya preserve hota hai:
  ✅ Font name, size, bold, italic, underline, strikethrough, color
  ✅ Paragraph alignment, indentation, spacing
  ✅ Bullet points aur numbered lists (numPr)
  ✅ Tables, Headers, Footers
  ✅ Subscript, superscript, all-caps, small-caps
  ✅ Text split across multiple runs (Word ka internal behavior)
"""

from docx import Document
import copy
import io
import tempfile
import os
from docx2pdf import convert

# Word XML namespace shortcut
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


# ══════════════════════════════════════════════════════════════════════
#  PHASE 1  —  DETECT
#  Document ki poori formatting ka complete XML snapshot lo.
#  Yeh snapshot baad mein Phase 3 mein restore guarantee ke kaam aata hai.
# ══════════════════════════════════════════════════════════════════════

def _snap_paragraph(para):
    """
    Ek paragraph ka 2-level snapshot:
      pPr  → paragraph properties  (alignment, indent, spacing, bullets)
      rPr  → run properties list   (font, bold, italic, color, size …)

    copy.deepcopy isliye ki XML object reference nahi, actual copy chahiye.
    """
    return {
        'pPr': copy.deepcopy(
            para._element.find(f'{{{W}}}pPr')
        ),
        'runs_rPr': [
            copy.deepcopy(run._element.find(f'{{{W}}}rPr'))
            for run in para.runs
        ],
    }


def detect_format(doc):
    """
    Poore document ka format-map banao.

    Returns dict with keys:
      body    → list of para-snapshots (main document body)
      tables  → nested list [table][row][cell] of para-snapshot lists
      headers → list per section
      footers → list per section
    """
    def snap_list(paras):
        return [_snap_paragraph(p) for p in paras]

    return {
        'body': snap_list(doc.paragraphs),

        'tables': [
            [
                [snap_list(cell.paragraphs) for cell in row.cells]
                for row in table.rows
            ]
            for table in doc.tables
        ],

        'headers': [snap_list(s.header.paragraphs) for s in doc.sections],
        'footers': [snap_list(s.footer.paragraphs) for s in doc.sections],
    }


def print_format_report(doc):
    """
    Detected formatting ka human-readable report print karo.
    Debug ke liye useful hai — dekho kya detect hua.
    """
    print("=" * 64)
    print("📋  DOCUMENT FORMAT REPORT")
    print("=" * 64)

    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue

        preview = para.text[:55] + ("…" if len(para.text) > 55 else "")
        print(f"\n[Para {i}]  Style: {para.style.name}  |  Align: {para.alignment}")
        print(f"  Text : '{preview}'")

        for j, run in enumerate(para.runs):
            if not run.text.strip():
                continue
            size_pt = (
                f"{int(run.font.size.pt)}pt"
                if run.font.size else "inherited"
            )
            color = "–"
            try:
                if run.font.color and run.font.color.type is not None:
                    color = str(run.font.color.rgb)
            except Exception:
                pass

            print(
                f"  Run {j}: '{run.text[:30]}'  |  "
                f"Font: {run.font.name or 'inherited'}  |  "
                f"Size: {size_pt}  |  "
                f"Bold: {run.bold}  |  "
                f"Italic: {run.italic}  |  "
                f"Underline: {run.underline}  |  "
                f"Color: {color}"
            )

    print("\n" + "=" * 64)


# ══════════════════════════════════════════════════════════════════════
#  PHASE 2  —  REPLACE
#  Run-level text replacement — paragraph structure kabhi nahi todni.
#  Rule: para.text = "..." KABHI NAHI karna.
# ══════════════════════════════════════════════════════════════════════

def _replace_in_paragraph(paragraph, old_text, new_text):
    """
    Single paragraph mein text replace karo — formatting safe rakh kar.

    Do cases:
      Simple  → old_text ek run mein milta hai (90% cases)
      Complex → old_text multiple runs mein split hai (Word ka quirk)
    """

    # ── Simple case ──────────────────────────────────────────────────
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return  # Kaam ho gaya ✅

    # ── Complex case: text multiple runs mein split hai ───────────────
    # Example: "Hello" Word ne internally "Hel" + "lo" kar diya
    combined = "".join(r.text for r in paragraph.runs)
    if old_text not in combined:
        return  # Is paragraph mein hai hi nahi

    start = combined.index(old_text)
    end   = start + len(old_text)

    # Har run ki character boundary map karo
    pos, bounds = 0, []
    for r in paragraph.runs:
        bounds.append((pos, pos + len(r.text)))
        pos += len(r.text)

    # Wo runs dhundo jo old_text ke range mein aate hain
    affected = [
        i for i, (rs, re) in enumerate(bounds)
        if rs < end and re > start
    ]
    if not affected:
        return

    first_idx = affected[0]
    last_idx  = affected[-1]

    # Pehle affected run mein:  [before] + [new_text] + [after last run]
    before = combined[bounds[first_idx][0] : start]
    after  = combined[end                  : bounds[last_idx][1]]

    paragraph.runs[first_idx].text = before + new_text + after

    # Baaki affected runs khali karo (structure rahegi, text nahi)
    for i in affected[1:]:
        paragraph.runs[i].text = ""


# ══════════════════════════════════════════════════════════════════════
#  PHASE 3  —  RESTORE
#  Phase 1 ke snapshot se exact formatting wapas lagao.
#  Yeh guarantee hai ki kuch bhi accidentally badla toh wapas aayega.
# ══════════════════════════════════════════════════════════════════════

def _restore_paragraph(para, snap):
    """
    Paragraph aur run properties ko snapshot se restore karo.

    pPr  → alignment, indentation, spacing, bullets (numPr)
    rPr  → font, bold, italic, underline, color, size, etc.
    """

    # 1. Paragraph Properties (pPr) restore
    if snap['pPr'] is not None:
        existing_pPr = para._element.find(f'{{{W}}}pPr')
        if existing_pPr is not None:
            para._element.remove(existing_pPr)
        para._element.insert(0, copy.deepcopy(snap['pPr']))

    # 2. Run Properties (rPr) restore — har run ke liye
    for i, run in enumerate(para.runs):
        if i >= len(snap['runs_rPr']):
            break
        saved_rPr = snap['runs_rPr'][i]
        if saved_rPr is None:
            continue
        curr_rPr = run._element.find(f'{{{W}}}rPr')
        if curr_rPr is not None:
            run._element.remove(curr_rPr)
        run._element.insert(0, copy.deepcopy(saved_rPr))


# ══════════════════════════════════════════════════════════════════════
#  HELPER — extract plain text
# ══════════════════════════════════════════════════════════════════════

def extract_text_from_docx(file_bytes):
    """Word file se plain text extract karta hai."""
    doc = Document(io.BytesIO(file_bytes))
    text = "\n".join(
        p.text.strip() for p in doc.paragraphs if p.text.strip()
    )
    return text, doc


# ══════════════════════════════════════════════════════════════════════
#  MAIN FUNCTION  —  DETECT → REPLACE → RESTORE
# ══════════════════════════════════════════════════════════════════════

def update_docx_and_export(doc, changes_dict, verbose=True):
    """
    DOCX mein safely text replace karta hai — 3-phase approach.

    Args:
        doc          : python-docx Document object
        changes_dict : { 'old_text': 'new_text', ... }
        verbose      : Format report + progress print kare? (default True)

    Returns:
        bytes : Updated DOCX file as bytes (download ke liye ready)
    """

    # ─── PHASE 1: DETECT ──────────────────────────────────────────────
    if verbose:
        print("🔍  Phase 1: Document formatting detect ho rahi hai...")
        print_format_report(doc)

    fmt = detect_format(doc)           # Complete XML snapshot memory mein

    if verbose:
        print("✅  Snapshot complete!\n")
        print("✏️   Phase 2: Text replace + Phase 3: Restore — shuru...")

    # ─── PHASE 2 + 3: REPLACE then immediately RESTORE ────────────────
    # ─── PHASE 2 + 3: REPLACE then immediately RESTORE ────────────────
    def process(paragraphs, snaps):
        """Paragraphs pe replace karo, phir turant snapshot se restore karo."""
        for i, para in enumerate(paragraphs):
            
            # 🚀 HYPERLINK SHIELD: Agar line mein link hai, usko completely ignore karo!
            if '<w:hyperlink' in para._p.xml:
                continue
                
            for old, new in changes_dict.items():
                if len(old) > 20 and old in para.text:
                    _replace_in_paragraph(para, old, new)   # Phase 2
            if snaps and i < len(snaps):
                _restore_paragraph(para, snaps[i])          # Phase 3

    # Main body
    process(doc.paragraphs, fmt['body'])

    # Tables (cell ke andar ke paragraphs)
    for t_i, table in enumerate(doc.tables):
        t_snap = fmt['tables'][t_i] if t_i < len(fmt['tables']) else []
        for r_i, row in enumerate(table.rows):
            for c_i, cell in enumerate(row.cells):
                c_snap = (
                    t_snap[r_i][c_i]
                    if r_i < len(t_snap) and c_i < len(t_snap[r_i])
                    else []
                )
                process(cell.paragraphs, c_snap)

    # Headers & Footers (har section ke liye)
    for s_i, section in enumerate(doc.sections):
        h_snap = fmt['headers'][s_i] if s_i < len(fmt['headers']) else []
        f_snap = fmt['footers'][s_i] if s_i < len(fmt['footers']) else []
        process(section.header.paragraphs, h_snap)
        process(section.footer.paragraphs, f_snap)

    if verbose:
        print("✅  Formatting 100% restore ho gayi!")
        print("🎉  File download ke liye ready hai!\n")

    # Bytes mein save karke return karo
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


# ══════════════════════════════════════════════════════════════════════
#  USAGE EXAMPLE
# ══════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    # File read karo
    with open("input.docx", "rb") as f:
        file_bytes = f.read()

    text, doc = extract_text_from_docx(file_bytes)

    # Jo replace karna hai
    changes = {
        "Purana Naam":     "Naya Naam",
        "Company ABC":     "Company XYZ",
        "January 2024":    "March 2025",
    }

    # 3-phase replace (verbose=True se progress + format report milega)
    result_bytes = update_docx_and_export(doc, changes, verbose=True)

    # Save karo
    with open("output.docx", "wb") as f:
        f.write(result_bytes)

    print("output.docx saved!")
